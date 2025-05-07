import re, docx, nltk, pandas as pd, zipfile, os, io, difflib, fitz, pytesseract
from tempfile import TemporaryDirectory
from pdf2image import convert_from_path
from django.core.files.base import ContentFile
from docx import Document
from nltk.corpus import stopwords
from collections import Counter
from django.db import transaction
from .models import Palabras, ConteoTotal, Provincia, Empresa, Reporte

nltk.download('stopwords')

def count_frequent_words(doc_paths):
    # Inicializar contador global
    total_word_counts = Counter()

    # Iterar sobre cada documento
    for doc_path in doc_paths:
        # Verificar extensión del archivo
        if doc_path.lower().endswith('.docx'):
            # Leer documento Word
            doc = docx.Document(doc_path)

            # Extraer texto del documento
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + " "

        elif doc_path.lower().endswith('.txt'):
            # Leer archivo de texto
            with open(doc_path, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            raise ValueError(f"Formato de archivo no soportado para {doc_path}. Use .docx o .txt")

        # Convertir a minúsculas y eliminar caracteres especiales (excepto letras)
        text = text.lower()
        text = re.sub(r'[^\w\s]', '', text)

        # Filtrar las palabras que contengan solo letras (eliminar números)
        text = re.sub(r'\b\d+\b', '', text)

        # Obtener lista de stop words en español
        spanish_stop_words = set(stopwords.words('spanish'))

        # Dividir el texto en palabras
        words = text.split()

        # Filtrar stop words y palabras vacías
        meaningful_words = [word for word in words if word not in spanish_stop_words and word.isalpha()]

        # Actualizar el contador global con las palabras del documento actual
        total_word_counts.update(meaningful_words)

    # Convertir a diccionario con las palabras más comunes
    most_common_dict = dict(total_word_counts.most_common())

    return most_common_dict


def guardar_conteo_en_bd(reporte, word_counts):
    """
    Guarda el conteo de palabras en la base de datos, acumulando
    para palabras del mismo año (sin importar la empresa).
    """
    with transaction.atomic():
        for palabra, cantidad in word_counts.items():
            palabra_obj, _ = Palabras.objects.get_or_create(descripcion=palabra)

            # Buscar si ya existe conteo para esta palabra en este año
            conteos_previos = ConteoTotal.objects.filter(
                reporte__anio=reporte.anio,
                palabra=palabra_obj
            )

            if conteos_previos.exists():
                # Sumamos al primer conteo encontrado
                conteo = conteos_previos.first()
                conteo.cantidad += cantidad
                conteo.save()
            else:
                # Si no existe, creamos uno nuevo vinculado al reporte actual
                ConteoTotal.objects.create(
                    reporte=reporte,
                    palabra=palabra_obj,
                    cantidad=cantidad
                )


def insertar_provincias(archivo_excel):
    # Leer el archivo Excel
    df = pd.read_excel(archivo_excel)

    # Insertar las provincias en la base de datos
    for _, row in df.iterrows():
        nombre_provincia = row['provincia']
        # Verificar si la provincia ya existe, si no, crearla
        if not Provincia.objects.filter(nombre=nombre_provincia).exists():
            Provincia.objects.create(nombre=nombre_provincia)

    print("Provincias importadas correctamente.")


def insertar_empresas(archivo_excel):
    # Leer el archivo Excel
    df = pd.read_excel(archivo_excel)

    for _, row in df.iterrows():
        nombre_empresa = row['NOMBRE DE LA ENTIDAD']
        ruc_empresa = row['IDENTIFICACIÓN']
        nombre_provincia = row['provincia']

        # Obtener o crear la provincia
        provincia, _ = Provincia.objects.get_or_create(nombre=nombre_provincia)

        # Verificar si la empresa ya existe por nombre o por ruc
        if not Empresa.objects.filter(nombre=nombre_empresa).exists() and not Empresa.objects.filter(ruc=ruc_empresa).exists():
            Empresa.objects.create(
                nombre=nombre_empresa,
                ruc=ruc_empresa,
                provincia=provincia
            )
        else:
            print(f"Empresa '{nombre_empresa}' con RUC '{ruc_empresa}' ya existe. No se insertó.")

    print("Empresas importadas correctamente.")


def procesar_zip_reportes(zip_file):
    """
    Procesa un archivo ZIP que contiene .pdf, .docx o .txt, crea instancias de Reporte y las guarda.
    """
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        with TemporaryDirectory() as temp_dir:
            for filename in zip_ref.namelist():
                if not filename.endswith(('.pdf', '.docx', '.txt')):
                    continue

                file_data = zip_ref.read(filename)

                # Guardar temporalmente el archivo para procesarlo si es necesario
                temp_path = os.path.join(temp_dir, os.path.basename(filename))
                with open(temp_path, 'wb') as f:
                    f.write(file_data)

                # Si es PDF, convertirlo a .docx
                if filename.endswith('.pdf'):
                    try:
                        temp_path = pdf_to_docx(temp_path, temp_dir)
                    except Exception as e:
                        print(f"Error al convertir PDF {filename} a DOCX: {e}")
                        continue  # Saltar archivo si falla la conversión

                # Leer texto
                texto = ''
                if temp_path.endswith('.docx'):
                    try:
                        doc = Document(temp_path)
                        texto = '\n'.join([p.text for p in doc.paragraphs])
                    except Exception as e:
                        print(f"Error al leer DOCX {filename}: {e}")
                        continue
                elif temp_path.endswith('.txt'):
                    with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                        texto = f.read()

                # Detectar año
                match_anio = re.search(r'\b(20\d{2}|19\d{2})\b', texto)
                anio = int(match_anio.group()) if match_anio else 0

                # Detectar empresa
                empresa_asignada = Empresa.objects.filter(nombre__iexact='Desconocido').first()
                umbral_similitud = 0.8
                for empresa in Empresa.objects.exclude(nombre__iexact='Desconocido'):
                    nombre_empresa = empresa.nombre.lower()
                    texto_busqueda = texto.lower()
                    for palabra in texto_busqueda.split():
                        similitud = difflib.SequenceMatcher(None, nombre_empresa, palabra).ratio()
                        if similitud >= umbral_similitud:
                            empresa_asignada = empresa
                            break
                    else:
                        for i in range(len(texto_busqueda.split()) - len(nombre_empresa.split()) + 1):
                            fragmento = " ".join(texto_busqueda.split()[i:i + len(nombre_empresa.split())])
                            similitud = difflib.SequenceMatcher(None, nombre_empresa, fragmento).ratio()
                            if similitud >= umbral_similitud:
                                empresa_asignada = empresa
                                break
                    if empresa_asignada and empresa_asignada.nombre != 'Desconocido':
                        break

                # Guardar reporte en la base de datos
                reporte = Reporte(empresa=empresa_asignada, anio=anio)
                reporte.archivo.save(os.path.basename(filename), ContentFile(file_data))
                reporte.save()



def pdf_to_docx(pdf_path, output_dir):
    # Extraer el nombre del archivo sin la extensión .pdf
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    # Crear la ruta de salida del archivo .docx
    output_docx_path = os.path.join(output_dir, f"{pdf_name}.docx")
    
    doc = Document()

    with fitz.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf):
            text = page.get_text().strip()
            if text:
                doc.add_paragraph(f"[Página {page_num + 1}]")
                doc.add_paragraph(text)
            else:
                with TemporaryDirectory() as tmpdir:
                    images = convert_from_path(pdf_path, first_page=page_num+1, last_page=page_num+1, output_folder=tmpdir)
                    for image in images:
                        ocr_text = pytesseract.image_to_string(image, lang='spa')
                        doc.add_paragraph(f"[Página {page_num + 1} - OCR]")
                        doc.add_paragraph(ocr_text.strip())

    # Guardar el archivo .docx generado
    doc.save(output_docx_path)
    return output_docx_path


# Ejemplo de uso:
#pdf_to_docx(
#    r'C:\Users\USUARIO\Downloads\drive-download-20250505T203410Z-1-001\EMPRESA 258 PLASTICOS RIVAL CIA LTDA\informe 2021.pdf',
#    r'C:\Users\USUARIO\Downloads\Prácticas\Programa\Contpal\media\Docxs'
#)
#insertar_provincias(r'C:\Users\USUARIO\Downloads\Empresas.xlsx')
#insertar_empresas(r'C:\Users\USUARIO\Downloads\Empresas.xlsx')
